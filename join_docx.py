
from docx import Document

def add_paragraph_with_formatting(destination_doc, paragraph):
    """Добавить абзац вместе с его стилевым форматированием."""
    new_paragraph = destination_doc.add_paragraph()
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.color.rgb = run.font.color.rgb
        new_run.style.name = run.style.name

def add_table(destination_doc, table):
    """Добавить таблицу с сохранением форматирования."""
    new_table = destination_doc.add_table(rows=0, cols=table.columns)
    for row in table.rows:
        cells = [cell.text for cell in row.cells]
        new_row = new_table.add_row().cells
        for idx, cell in enumerate(cells):
            new_row[idx].text = cell

# Путь к файлу результирующего документа
output_doc_path = 'merged_document.docx'

# Создаем новый документ, в который будем добавлять содержимое других документов
merged_document = Document()

# Список DOCX файлов для объединения
# Список DOCX-файлов для объединения
docx_files = ['/Applications/programming/kg_llm/dozor/dozor.docx',
              '/Applications/programming/kg_llm/dozor/install.docx',
              '/Applications/programming/kg_llm/dozor/usage.docx']

for docx_file in docx_files:
    # Открываем текущий DOCX-файл
    sub_doc = Document(docx_file)
    
    # Добавляем содержимое документа в результирующий файл
    for element in sub_doc.element.body:
        if element.tag.endswith('p'):
            # Копируем абзац
            para = sub_doc.paragraphs[element]
            add_paragraph_with_formatting(merged_document, para)

# Сохраняем объединенный документ
merged_document.save(output_doc_path)

print("DOCX-файлы были успешно объединены с сохранением базового форматирования!")